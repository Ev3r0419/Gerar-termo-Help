from docx import Document
import os
import sys
from datetime import datetime
from docx2pdf import convert
from pathlib import Path


def get_resource_path(relative_path):
    """Permite acessar arquivos internos, tanto no modo script quanto empacotado no .exe"""
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller armazena arquivos dentro do diretório temporário _MEIPASS
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)   


class GeradorDeTermos:

    #Gera termo de empréstimo de equipamentos
    def preencher_termo_equipamento(self, nome, cpf, setor, empresa, equipamento, marca, serie, patrimonio, estadoequip, tecnico):
        """Gera termo de empréstimo de equipamento."""
        documento = Document(get_resource_path("Termo de Responsabilidade de Emprestimo de Equipamento.docx"))

        referencias = {
            "NNNN": nome,
            "CCCC": cpf,
            "SSSS": setor,
            "EEEE": empresa,
            "DDDD": equipamento,
            "MMMM": marca,
            "XXXX": serie,
            "PPPP": patrimonio,
            "ZZZZ": estadoequip,
            "TTTT": tecnico,  
            "DD": str(datetime.now().day),
            "MM": str(datetime.now().month),
            "AAAA": str(datetime.now().year)
        }

        self._substituir_textos(documento, referencias)
        nome_arquivo = f"Equipamentos - {nome}.docx"
        documento.save(nome_arquivo)
        convert(nome_arquivo)

        return Path(f"Equipamentos - {nome}.pdf")
    

    #Gera termo de empréstimo de equipamento do Telecom.
    def preencher_termo_telecom(self, nome, cpf, setor, empresa,equipamento, marca, serie, numero, tecnico):
        documento = Document(get_resource_path("Termo Telecom.docx"))

        referencias = {
            "NNNN": nome,
            "CCCC": cpf,
            "SSSS": setor,
            "EEEE": empresa,
            "DDDD": equipamento,
            "MMMM": marca,
            "XXXX": serie,
            "LLLL": numero,
            "TTTT": tecnico,  
            "DD": str(datetime.now().day),
            "MM": str(datetime.now().month),
            "AAAA": str(datetime.now().year)
        }

        self._substituir_textos(documento, referencias)
        nome_arquivo = f"Equipamentos Telecom - {nome}.docx"
        documento.save(nome_arquivo)
        convert(nome_arquivo)

        return Path(f"Equipamentos Telecom - {nome}.pdf")

    #
    def preencher_termo_vpn(self, nome, cargo, departamento):
        """Gera termo de VPN."""
        documento = Document(get_resource_path("Modelo de termo VPN.docx"))

        referencias = {
            "NNNN": nome,
            "CG": cargo,
            "DPT": departamento,
            "DD": str(datetime.now().day),
            "MM": str(datetime.now().month),
            "AAAA": str(datetime.now().year)
        }

        self._substituir_textos(documento, referencias)
        nome_arquivo = f"VPN - {nome}.docx"
        documento.save(nome_arquivo)
        convert(nome_arquivo)

        return Path(f"VPN - {nome}.pdf")

    def preencher_termo_adm(self, nome, cpf):
        """Gera termo de administrador local."""
        documento = Document(get_resource_path("Modelo do Termo de Administrador Local.docx"))

        referencias = {
            "NNNN": nome,
            "CCCC": cpf,
            "DD": str(datetime.now().day),
            "MM": str(datetime.now().month),
            "AAAA": str(datetime.now().year)
        }

        self._substituir_textos(documento, referencias)
        nome_arquivo = f"ADM - {nome}.docx"
        documento.save(nome_arquivo)
        convert(nome_arquivo)

        return Path(f"ADM - {nome}.pdf")

    # ============================================================
    # 🔧 Função auxiliar — substitui o texto nos parágrafos e tabelas
    # ============================================================
    def _substituir_textos(self, documento, referencias):
        """Substitui textos no documento Word mantendo formatação."""
        for paragrafo in documento.paragraphs:
            for run in paragrafo.runs:
                for codigo, valor in referencias.items():
                    if codigo in run.text:
                        run.text = run.text.replace(codigo, valor)

        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            for codigo, valor in referencias.items():
                                if codigo in run.text:
                                    run.text = run.text.replace(codigo, valor)
